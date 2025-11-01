# backend/docx_parser.py
import re
from collections import defaultdict
from docx import Document

# Patterns
BRACKETED_GENERIC = re.compile(r"\[\s*_{2,}\s*\]")                   # [_________]
BRACKETED_NAMED   = re.compile(r"\[[^\[\]\r\n]{1,60}\]")            # [Company Name], [Date of Safe], etc.

# Heuristic: tokens that likely name money / dates / parties
MONEY_HINT_TOKENS = ["purchase amount", "purchase price", "price", "amount", "consideration", "principal", "valuation cap", "cap"]
DATE_HINT_TOKENS  = ["date", "effective date", "closing date", "date of safe"]
COMPANY_HINT_TOKENS = ["company", "issuer", "corporation", "startup", "llc", "inc"]
INVESTOR_HINT_TOKENS = ["investor", "purchaser", "buyer", "lender", "holder"]

QUOTE_PAT  = re.compile(r"[“\"]([^”\"]+)[”\"]")   # capture last quoted phrase “like this” or "like this"

def _titleish_phrases_around(text_before: str, text_after: str) -> str | None:
    """
    Try to find a meaningful label near the placeholder.
    Priority:
      1) Last quoted phrase before (e.g., “Purchase Amount”)
      2) Capitalized phrase nearby that contains key tokens
      3) Quoted or capitalized phrase right after
    """
    # Clean weird trailing punctuation
    text_before = re.sub(r"[\$\(\)\[\]]+$", "", text_before.strip())

    # 1️⃣ Last quoted phrase before placeholder
    for m in reversed(list(QUOTE_PAT.finditer(text_before))):
        phrase = m.group(1).strip()
        if 2 <= len(phrase) <= 60 and not any(ch in phrase for ch in "[]"):
            return phrase

    # 2️⃣ Capitalized or key-token phrase before
    tail = text_before[-160:]  # widen context
    tail = re.sub(r"[\[\]\(\)\$]+", "", tail)
    chunks = re.split(r"[.;:\n]", tail)
    if chunks:
        cand = chunks[-1].strip()
        l = cand.lower()
        if any(t in l for t in MONEY_HINT_TOKENS + DATE_HINT_TOKENS + COMPANY_HINT_TOKENS + INVESTOR_HINT_TOKENS):
            cand = re.sub(r"\s+", " ", cand).strip(" -:")
            # Normalize casing
            words = cand.split()
            if len(words) > 1:
                cand = " ".join(w.capitalize() for w in words)
            else:
                cand = words[0].capitalize()
            return cand

    # 3️⃣ Check for label right after
    head = text_after[:120]
    head = re.sub(r"[\[\]\(\)\$]+", "", head)
    for qm in QUOTE_PAT.finditer(head):
        phrase = qm.group(1).strip()
        if 2 <= len(phrase) <= 60 and not any(ch in phrase for ch in "[]"):
            return phrase

    # 4️⃣ If pattern like “the Purchase Amount” appears nearby
    combo = (text_before[-200:] + text_after[:200]).lower()
    for token in MONEY_HINT_TOKENS + DATE_HINT_TOKENS:
        if token in combo:
            return token.title()

    return None

def _rename_generic_placeholder_in_text(text: str) -> tuple[str, list[str]]:
    """
    Replace generic [____] placeholders with context-based names when possible, otherwise enumerate.
    Returns (new_text, keys_found)
    """
    keys = []
    out = []
    i = 0
    dup_counts = defaultdict(int)

    # Find all matches (both named and generic) with spans
    matches = []
    for m in BRACKETED_GENERIC.finditer(text):
        matches.append(("GEN", m.span()))
    for m in BRACKETED_NAMED.finditer(text):
        # exclude the generic ones from this list
        frag = text[m.span()[0]:m.span()[1]]
        if BRACKETED_GENERIC.fullmatch(frag):
            continue
        matches.append(("NAM", m.span()))

    matches.sort(key=lambda t: t[1][0])

    last = 0
    for kind, (a, b) in matches:
        # Append text before match
        before = text[last:a]
        frag = text[a:b]
        last = b

        if kind == "NAM":
            # Already-named placeholder; keep as-is
            out.append(before)
            out.append(frag)
            keys.append(frag)
            continue

        # kind == GEN
        # Find a label near the placeholder
        label = _titleish_phrases_around(text[:a], text[b:])
        if label:
            key = f"[{label}]"
        else:
            key = "[Blank]"

        # ensure uniqueness
        dup_counts[key] += 1
        if dup_counts[key] > 1:
            key = f"{key}#{dup_counts[key]}"

        out.append(before)
        out.append(key)
        keys.append(key)

    # trailing text
    out.append(text[last:])
    return "".join(out), keys

def find_placeholders(docx_path: str) -> list[str]:
    """
    Load DOCX, rename generic placeholders to semantic keys when possible by
    using nearby context (quoted phrases like “Purchase Amount”), otherwise enumerate.
    Save the modified doc back to docx_path, and return unique placeholder keys in reading order.
    """
    doc = Document(docx_path)
    found_keys = []

    # We will rebuild paragraph text (formatting may be slightly simplified).
    for p in doc.paragraphs:
        text = p.text
        if "[" not in text or "]" not in text:
            continue
        new_text, keys = _rename_generic_placeholder_in_text(text)
        if keys:
            # replace paragraph runs with single run
            p.clear()
            run = p.add_run(new_text)
            found_keys.extend(keys)

    # Also handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    if "[" not in text or "]" not in text:
                        continue
                    new_text, keys = _rename_generic_placeholder_in_text(text)
                    if keys:
                        p.clear()
                        p.add_run(new_text)
                        found_keys.extend(keys)

    # Save patched doc (so fill_placeholders can replace by new keys)
    doc.save(docx_path)

    # Return unique keys in order of first appearance
    unique = []
    seen = set()
    for k in found_keys:
        if k not in seen:
            seen.add(k)
            unique.append(k)
    return unique

def fill_placeholders(original_docx_path: str, working_docx_path: str, mapping: dict[str, str]):
    """
    Simple text replace in the working docx for each [Key] -> value.
    We operate on the working copy in paragraphs + tables.
    """
    doc = Document(working_docx_path)

    def apply_on_text(t: str) -> str:
        for k, v in mapping.items():
            t = t.replace(k, str(v))
        return t

    for p in doc.paragraphs:
        txt = p.text
        if "[" in txt and "]" in txt:
            p.clear()
            p.add_run(apply_on_text(txt))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text
                    if "[" in txt and "]" in txt:
                        p.clear()
                        p.add_run(apply_on_text(txt))

    doc.save(working_docx_path)
